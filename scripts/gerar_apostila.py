#!/usr/bin/env python3
"""
Gerador automático de apostilas EBE usando Gemini AI.
Gera apostilas de 15-20 páginas e faz upload directo para Google Drive.

Uso:
  python scripts/gerar_apostila.py              # Gerar apostilas
  python scripts/gerar_apostila.py --test-drive  # Testar conexão ao Drive
"""

import json
import os
import re
import sys
import time
import unicodedata
import traceback

def log(msg):
    print(msg, flush=True)

# ──────────────────────────────────────────────────────────────
# CONFIGURAÇÃO
# ──────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
DRIVE_FOLDER_ID = os.environ.get("GOOGLE_DRIVE_FOLDER_ID", "")
CREDENTIALS_JSON = os.environ.get("GOOGLE_DRIVE_CREDENTIALS", "")
MAPA_PATH = os.path.join(SCRIPT_DIR, "mapa_apostilas.json")
ESTADO_PATH = os.path.join(SCRIPT_DIR, "estado.json")
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "..", "output")
BATCH_SIZE = int(os.environ.get("BATCH_SIZE", "5"))
MODEL_NAME = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
DELAY_ENTRE_APOSTILAS = 90
MAX_UPLOAD_RETRIES = 3
UPLOAD_RETRY_DELAY = 5

from docx.shared import RGBColor, Pt, Cm
COR_PRIMARIA   = RGBColor(0x1B, 0x3A, 0x5C)
COR_SECUNDARIA = RGBColor(0x2E, 0x7D, 0x4F)
COR_TEXTO      = RGBColor(0x1A, 0x1A, 0x1A)
COR_CITACAO    = RGBColor(0x55, 0x55, 0x55)
HEX_PRIMARIA   = "1B3A5C"
HEX_SECUNDARIA = "2E7D4F"
FONTE          = "Garamond"
ASSETS_DIR = os.path.join(SCRIPT_DIR, "..", "src", "estilos", "_assets")
LOGO_PATH  = os.path.join(ASSETS_DIR, "logo_ebe.png")


# ──────────────────────────────────────────────────────────────
# GOOGLE DRIVE
# ──────────────────────────────────────────────────────────────
_drive_service = None

def validar_credenciais():
    """Valida se as credenciais do Drive são JSON válido."""
    if not CREDENTIALS_JSON:
        return False, "GOOGLE_DRIVE_CREDENTIALS vazio"
    try:
        creds = json.loads(CREDENTIALS_JSON)
        if "type" not in creds:
            return False, "JSON não contém 'type' (não é service account?)"
        if creds.get("type") != "service_account":
            return False, f"type='{creds.get('type')}' (esperado: service_account)"
        if "client_email" not in creds:
            return False, "JSON não contém 'client_email'"
        return True, f"OK — {creds.get('client_email', 'N/A')}"
    except json.JSONDecodeError as e:
        return False, f"JSON inválido: {e}"


def get_drive_service():
    global _drive_service
    if _drive_service:
        return _drive_service

    log("    🔑 A autenticar no Google Drive...")
    ok, msg = validar_credenciais()
    if not ok:
        raise RuntimeError(f"Credenciais inválidas: {msg}")

    from google.oauth2 import service_account
    from googleapiclient.discovery import build

    creds_info = json.loads(CREDENTIALS_JSON)
    credentials = service_account.Credentials.from_service_account_info(
        creds_info, scopes=["https://www.googleapis.com/auth/drive.file"]
    )
    _drive_service = build("drive", "v3", credentials=credentials)
    log("    ✅ Autenticado no Google Drive")
    return _drive_service


def test_drive_connection():
    """Testa a conexão ao Google Drive e à pasta mãe."""
    log("=" * 60)
    log("  TESTE DE CONEXÃO — Google Drive")
    log("=" * 60)

    # 1. Validar credenciais
    log("\n📋 1. Validando credenciais...")
    ok, msg = validar_credenciais()
    log(f"   {'✅' if ok else '❌'} {msg}")
    if not ok:
        return False

    # 2. Autenticar
    log("\n📋 2. Autenticando...")
    try:
        service = get_drive_service()
    except Exception as e:
        log(f"   ❌ Falha na autenticação: {e}")
        log(f"   📋 Detalhes: {traceback.format_exc()}")
        return False

    # 3. Verificar pasta mãe
    log(f"\n📋 3. Verificando pasta mãe: {DRIVE_FOLDER_ID}...")
    if not DRIVE_FOLDER_ID:
        log("   ❌ GOOGLE_DRIVE_FOLDER_ID vazio")
        return False

    try:
        folder = service.files().get(
            fileId=DRIVE_FOLDER_ID,
            fields="id, name, mimeType"
        ).execute()
        log(f"   ✅ Pasta encontrada: '{folder.get('name')}' (ID: {folder.get('id')})")
    except Exception as e:
        log(f"   ❌ Pasta não encontrada ou sem acesso: {e}")
        log(f"   💡 Verifique se a service account tem acesso à pasta {DRIVE_FOLDER_ID}")
        log(f"   📋 Detalhes: {traceback.format_exc()}")
        return False

    # 4. Testar criação de subpasta
    log("\n📋 4. Testando criação de subpasta...")
    try:
        test_folder = service.files().create(
            body={
                'name': '_teste_conexao_EBE',
                'mimeType': 'application/vnd.google-apps.folder',
                'parents': [DRIVE_FOLDER_ID]
            },
            fields='id, name'
        ).execute()
        log(f"   ✅ Subpasta criada: '{test_folder['name']}' (ID: {test_folder['id']})")

        # Limpar subpasta de teste
        service.files().delete(fileId=test_folder['id']).execute()
        log(f"   🧹 Subpasta de teste removida")
    except Exception as e:
        log(f"   ❌ Falha ao criar subpasta: {e}")
        log(f"   📋 Detalhes: {traceback.format_exc()}")
        return False

    # 5. Testar upload de ficheiro pequeno
    log("\n📋 5. Testando upload de ficheiro...")
    try:
        import tempfile
        from googleapiclient.http import MediaInMemoryUpload

        test_content = b"Teste de conexao EBE - pode ser apagado"
        media = MediaInMemoryUpload(test_content, mimetype='text/plain')
        test_file = service.files().create(
            body={'name': '_teste_upload_EBE.txt', 'parents': [DRIVE_FOLDER_ID]},
            media_body=media, fields='id, name'
        ).execute()
        log(f"   ✅ Ficheiro enviado: '{test_file['name']}' (ID: {test_file['id']})")

        # Limpar
        service.files().delete(fileId=test_file['id']).execute()
        log(f"   🧹 Ficheiro de teste removido")
    except Exception as e:
        log(f"   ❌ Falha no upload: {e}")
        log(f"   📋 Detalhes: {traceback.format_exc()}")
        return False

    log(f"\n{'=' * 60}")
    log("  ✅ TODOS OS TESTES PASSARAM — Drive pronto para upload")
    log(f"{'=' * 60}")
    return True


def sanitizar_nome(nome):
    for c in ['"', "'", '/', '\\', '|', ':', '*', '?', '<', '>']:
        nome = nome.replace(c, '')
    return nome.strip()[:100]


_pasta_cache = {}

def encontrar_ou_criar_pasta(service, nome, parent_id):
    nome = sanitizar_nome(nome)
    query = (
        f"name = '{nome}' and '{parent_id}' in parents and "
        f"mimeType = 'application/vnd.google-apps.folder' and trashed = false"
    )
    results = service.files().list(
        q=query, spaces='drive', fields='files(id, name)', pageSize=10
    ).execute()
    files = results.get('files', [])
    if files:
        log(f"    📁 Pasta existe: {nome}")
        return files[0]['id']
    log(f"    📁 Criando pasta: {nome}")
    folder = service.files().create(
        body={'name': nome, 'mimeType': 'application/vnd.google-apps.folder', 'parents': [parent_id]},
        fields='id'
    ).execute()
    return folder['id']


def garantir_caminho(service, parent_id, partes):
    current_id = parent_id
    for parte in partes:
        if parte:
            current_id = encontrar_ou_criar_pasta(service, parte, current_id)
    return current_id


def upload_para_drive(file_path, meta):
    """Faz upload de um ficheiro .docx para o Google Drive com retry logic."""
    log(f"    📂 Ficheiro: {file_path}")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Ficheiro não encontrado: {file_path}")

    size_kb = os.path.getsize(file_path) / 1024
    log(f"    📂 Tamanho: {size_kb:.0f} KB")

    if not DRIVE_FOLDER_ID:
        raise RuntimeError("GOOGLE_DRIVE_FOLDER_ID não definido")

    from googleapiclient.http import MediaFileUpload
    service = get_drive_service()

    partes = [
        sanitizar_nome(meta.get("instituto", "Outros")),
        sanitizar_nome(meta.get("escola", "Geral")),
        sanitizar_nome(meta.get("curso", "Curso")),
        sanitizar_nome(f"Módulo {meta.get('modulo', '1')}"),
    ]
    log(f"    📂 Destino: {' / '.join(partes)}")

    cache_key = " > ".join(partes)
    if cache_key in _pasta_cache:
        folder_id = _pasta_cache[cache_key]
        log(f"    📂 Folder ID (cache): {folder_id}")
    else:
        folder_id = garantir_caminho(service, DRIVE_FOLDER_ID, partes)
        _pasta_cache[cache_key] = folder_id
        log(f"    📂 Folder ID: {folder_id}")

    fname = os.path.basename(file_path)
    log(f"    📤 A enviar: {fname} → folder {folder_id}")

    # Retry logic para upload
    for attempt in range(1, MAX_UPLOAD_RETRIES + 1):
        try:
            media = MediaFileUpload(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                resumable=True
            )
            file = service.files().create(
                body={'name': fname, 'parents': [folder_id]},
                media_body=media, fields='id, name, webViewLink'
            ).execute()

            log(f"    ✅ Upload OK! ID: {file['id']}")
            return file
        except Exception as e:
            log(f"    ⚠️  Tentativa {attempt}/{MAX_UPLOAD_RETRIES} falhou: {type(e).__name__}: {e}")
            if attempt < MAX_UPLOAD_RETRIES:
                log(f"    ⏳ Aguardando {UPLOAD_RETRY_DELAY}s antes de tentar novamente...")
                time.sleep(UPLOAD_RETRY_DELAY)
            else:
                raise


# ──────────────────────────────────────────────────────────────
# DOCX
# ──────────────────────────────────────────────────────────────
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _add_horizontal_line(paragraph, color=HEX_PRIMARIA, size=8):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)

def configurar_estilos_base(doc):
    style = doc.styles["Normal"]
    style.font.name = FONTE
    style.font.size = Pt(12)
    style.font.color.rgb = COR_TEXTO
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.4
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width  = Cm(21.0)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

def inserir_logo(doc, caminho=None, largura_cm=5.5):
    caminho = caminho or LOGO_PATH
    if not os.path.exists(caminho):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(caminho, width=Cm(largura_cm))

def add_texto(doc, texto, font_size=12, bold=False, italic=False,
              color=None, alignment=None, indent_cm=0):
    p = doc.add_paragraph()
    if alignment: p.alignment = alignment
    if indent_cm: p.paragraph_format.left_indent = Cm(indent_cm)
    r = p.add_run(texto)
    r.font.name = FONTE; r.font.size = Pt(font_size)
    r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    return p

def cabecalho_rodape(doc, titulo_doc, codigo_doc):
    for section in doc.sections:
        section.different_first_page_header_footer = True
        header = section.header
        ph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = ph.add_run(f"Escola Bíblica Epignósis  ·  {titulo_doc}")
        r.font.name = FONTE; r.font.size = Pt(9)
        r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA
        footer = section.footer
        pf = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pf.add_run(f"{codigo_doc}  ·  ")
        r.font.name = FONTE; r.font.size = Pt(9); r.font.color.rgb = COR_CITACAO
        run = pf.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.name = FONTE; run.font.size = Pt(9); run.font.color.rgb = COR_CITACAO

def slugify(texto):
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("ascii")
    texto = re.sub(r'[^\w\s-]', '', texto).strip().lower()
    texto = re.sub(r'[\s_]+', '_', texto)
    return texto[:80]

def construir_docx(meta, conteudo_md, output_path):
    doc = Document()
    configurar_estilos_base(doc)
    codigo = meta.get("codigo", "EBE-APO-XXXX")
    cabecalho_rodape(doc, meta.get("titulo", "Apostila"), codigo)

    # CAPA
    doc.add_paragraph()
    inserir_logo(doc, LOGO_PATH, largura_cm=5.5)
    add_texto(doc, "ἐπίγνωσις  ·  Conhecer a Deus. Viver a Palavra. Manifestar o Reino.",
              font_size=11, italic=True, color=COR_SECUNDARIA, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    _add_horizontal_line(p, color=HEX_SECUNDARIA, size=8)
    for _ in range(2): doc.add_paragraph()
    if meta.get("supratitulo"):
        add_texto(doc, meta["supratitulo"].upper(), font_size=11, bold=True,
                  color=COR_SECUNDARIA, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_texto(doc, meta.get("titulo", ""), font_size=26, bold=True,
              color=COR_PRIMARIA, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if meta.get("subtitulo"):
        add_texto(doc, meta["subtitulo"], font_size=13, italic=True,
                  color=COR_TEXTO, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph()
    _add_horizontal_line(p, color=HEX_SECUNDARIA, size=4)
    add_texto(doc, meta.get("info_linha", f"Código {codigo} · 2026"),
              font_size=10, color=COR_CITACAO, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    # MARCO FILOSÓFICO
    for _ in range(5): doc.add_paragraph()
    add_texto(doc, "MARCO FILOSÓFICO", font_size=12, bold=True,
              color=COR_SECUNDARIA, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    _add_horizontal_line(p, color=HEX_SECUNDARIA, size=4)
    doc.add_paragraph()
    add_texto(doc,
        "\u201cAcreditamos que o verdadeiro conhecimento de Deus transforma "
        "a mente pela verdade das Escrituras, o coração pela acção do "
        "Espírito Santo e a vida pelo compromisso de viver e anunciar "
        "o Evangelho de Jesus Cristo.\u201d",
        font_size=14, italic=True, color=COR_PRIMARIA,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=2)
    doc.add_paragraph()
    add_texto(doc, "— Escola Bíblica Epignósis —", font_size=10,
              color=COR_CITACAO, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    # CONTEÚDO
    for linha in conteudo_md.strip().split('\n'):
        linha = linha.rstrip()
        if not linha.strip(): continue
        if linha.startswith('## '):
            texto = linha[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(texto)
            r.font.name = FONTE; r.font.size = Pt(13)
            r.font.bold = True; r.font.color.rgb = COR_PRIMARIA
        elif linha.startswith('### '):
            texto = linha[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(texto)
            r.font.name = FONTE; r.font.size = Pt(11.5)
            r.font.bold = True; r.font.color.rgb = COR_SECUNDARIA
        elif linha.startswith('> '):
            texto = linha[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.5)
            r = p.add_run(texto)
            r.font.name = FONTE; r.font.size = Pt(11)
            r.font.italic = True; r.font.color.rgb = COR_CITACAO
        elif linha.startswith('- ') or linha.startswith('* '):
            texto = re.sub(r'\*\*(.+?)\*\*', r'\1', linha[2:].strip())
            texto = re.sub(r'\*(.+?)\*', r'\1', texto)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            r = p.add_run("•  ")
            r.font.name = FONTE; r.font.size = Pt(12)
            r.font.bold = True; r.font.color.rgb = COR_SECUNDARIA
            r2 = p.add_run(texto)
            r2.font.name = FONTE; r2.font.size = Pt(12)
        elif linha.startswith('# '):
            texto = linha[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            r = p.add_run(texto.upper())
            r.font.name = FONTE; r.font.size = Pt(15)
            r.font.bold = True; r.font.color.rgb = COR_PRIMARIA
            _add_horizontal_line(p, color=HEX_PRIMARIA, size=6)
        else:
            texto = re.sub(r'\*\*(.+?)\*\*', r'\1', linha.strip())
            texto = re.sub(r'\*(.+?)\*', r'\1', texto)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(texto)
            r.font.name = FONTE; r.font.size = Pt(12)

    # SELO FINAL
    doc.add_paragraph()
    p = doc.add_paragraph()
    _add_horizontal_line(p, color=HEX_SECUNDARIA, size=4)
    doc.add_paragraph()
    add_texto(doc, "ESCOLA BÍBLICA EPIGNÓSIS", font_size=11, bold=True,
              color=COR_PRIMARIA, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_texto(doc, "Conhecer a Deus. Viver a Palavra. Manifestar o Reino.",
              font_size=10, italic=True, color=COR_SECUNDARIA, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_texto(doc, "Soli Deo Gloria", font_size=9, italic=True,
              color=COR_CITACAO, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.save(output_path)


# ──────────────────────────────────────────────────────────────
# GEMINI
# ──────────────────────────────────────────────────────────────
SYSTEM_PROMPT = (
    "Você é um teólogo, pedagogo e editor cristão sénior da Escola Bíblica Epignósis (EBE). "
    "Redija conteúdo académico cristão original em português europeu/Angola (pt-PT). "
    "Use referências bíblicas mas PARAfraseie o texto — não cite versículos longos palavra por palavra. "
    "Para referências curtas (até 10 palavras), cite directamente. "
    "Para passagens longas, resuma e indique a referência. "
    "Estilo académico formal, acessível e devocional. "
    "Use termos gregos/hebraicos quando relevante (com transliteração). "
    "Escreva em Markdown."
)


def gerar_conteudo_gemini(meta):
    import google.generativeai as genai
    log(f"  📡 Gemini API ({MODEL_NAME})...")
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel(model_name=MODEL_NAME, system_instruction=SYSTEM_PROMPT)

    prompt = (
        f"Escreva o conteúdo completo de uma apostila académica cristã.\n\n"
        f"CONTEXTO:\n"
        f"- Instituto: {meta['instituto']}\n"
        f"- Escola: {meta['escola']}\n"
        f"- Curso: {meta['curso']} ({meta.get('carga_horaria', 'N/D')})\n"
        f"- Módulo: {meta['modulo']}\n"
        f"- Apostila: {meta['titulo']}\n"
        f"- Código: {meta['codigo']}\n\n"
        f"ESTRUTURA OBRIGATÓRIA:\n"
        f"## FICHA TÉCNICA\n## SUMÁRIO\n## APRESENTAÇÃO DA APOSTILA\n"
        f"## OBJECTIVOS DE APRENDIZAGEM\n## VERSÍCULO-CHAVE\n## TEXTO-BASE PARA LEITURA\n"
        f"## 1. INTRODUÇÃO\n## 2. DESENVOLVIMENTO DO CONCEITO CENTRAL\n"
        f"### 2.1 Fundamentos bíblicos\n### 2.2 Desenvolvimento temático\n"
        f"### 2.3 Aprofundamento\n### 2.4 Dúvidas comuns\n### 2.5 Quadro de Destaque\n"
        f"## 3. APLICAÇÃO PRÁTICA\n## 4. SÍNTESE E CONCLUSÃO\n"
        f"## EXERCÍCIOS DE REVISÃO\n## ESTUDO BÍBLICO COMPLEMENTAR\n"
        f"## PARA A PRÓXIMA APOSTILA\n## GLOSSÁRIO\n## BIBLIOGRAFIA RECOMENDADA\n"
        f"## ANOTAÇÕES PESSOAIS\n"
    )

    log(f"  📤 Prompt: {len(prompt)} chars")
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            temperature=0.8, max_output_tokens=16384, top_p=0.95,
        ),
    )

    if not response:
        raise RuntimeError("Resposta vazia do Gemini")

    if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
        log(f"  ⚠️  Feedback: {response.prompt_feedback}")

    try:
        texto = response.text
    except Exception as e:
        log(f"  ❌ response.text erro: {e}")
        if hasattr(response, 'candidates') and response.candidates:
            for c in response.candidates:
                log(f"  Candidate finish_reason: {getattr(c, 'finish_reason', 'N/A')}")
        raise RuntimeError(f"Sem texto: {e}")

    if not texto or len(texto) < 100:
        raise RuntimeError(f"Conteúdo curto: {len(texto) if texto else 0} chars")

    texto = re.sub(r'^```(?:markdown)?\s*\n?', '', texto, flags=re.MULTILINE)
    texto = re.sub(r'\n?```\s*$', '', texto, flags=re.MULTILINE)
    log(f"  ✅ Resposta: {len(texto)} chars")
    return texto.strip()


# ──────────────────────────────────────────────────────────────
# ESTADO
# ──────────────────────────────────────────────────────────────
def carregar_estado():
    if os.path.exists(ESTADO_PATH):
        with open(ESTADO_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {"concluidas": [], "falhadas": [], "total_gerado": 0, "uploads": []}

def salvar_estado(estado):
    with open(ESTADO_PATH, 'w', encoding='utf-8') as f:
        json.dump(estado, f, ensure_ascii=False, indent=2)

def carregar_mapa():
    with open(MAPA_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)

def obter_proximas_apostilas(mapa, estado, batch_size):
    concluidas = set(estado.get("concluidas", []))
    pendentes = []
    for nivel in mapa["niveis"]:
        for inst in nivel["institutos"]:
            for escola in inst["escolas"]:
                for curso in escola["cursos"]:
                    carga = curso.get("carga_horaria", "N/D")
                    for modulo in curso["modulos"]:
                        for apo in modulo["apostilas"]:
                            if apo["id"] not in concluidas:
                                pendentes.append({
                                    "id": apo["id"],
                                    "titulo": apo["titulo"],
                                    "nivel": nivel["id"],
                                    "instituto": inst["nome"],
                                    "escola": escola["nome"],
                                    "curso": curso["nome"],
                                    "carga_horaria": carga,
                                    "modulo": f"{modulo['numero']} — {modulo['nome']}",
                                    "numero_apostila": apo["id"].split("-")[1].lstrip("0") or "1",
                                    "codigo": apo["id"],
                                })
    return pendentes[:batch_size]


# ──────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────
def main():
    # Modo de teste do Drive
    if "--test-drive" in sys.argv:
        ok = test_drive_connection()
        sys.exit(0 if ok else 1)

    log("=" * 60)
    log("  ESCOLA BÍBLICA EPIGNÓSIS — Gerador de Apostilas")
    log(f"  Modelo: {MODEL_NAME} | Batch: {BATCH_SIZE}")
    log("=" * 60)

    if not GEMINI_API_KEY:
        log("❌ GEMINI_API_KEY não definida!")
        sys.exit(1)
    log(f"✅ GEMINI_API_KEY: {len(GEMINI_API_KEY)} chars")

    # Diagnóstico detalhado do Drive
    drive_ok = False
    if DRIVE_FOLDER_ID and CREDENTIALS_JSON:
        ok_creds, msg_creds = validar_credenciais()
        log(f"{'✅' if ok_creds else '❌'} Credenciais: {msg_creds}")
        log(f"✅ DRIVE_FOLDER_ID: {DRIVE_FOLDER_ID[:12]}...")
        
        # Tentar autenticar mesmo que validação falhe
        if ok_creds:
            try:
                get_drive_service()
                drive_ok = True
                log(f"✅ Autenticação no Drive: sucesso")
            except Exception as e:
                log(f"⚠️  Autenticação no Drive falhou: {type(e).__name__}: {e}")
                log(f"💡 A gerar apostilas SEM upload para Drive")
                log(f"📋 Para diagnosticar: python scripts/gerar_apostila.py --test-drive")
        else:
            log(f"⚠️  Credenciais inválidas — a gerar SEM upload")
    else:
        log(f"⚠️  Drive não configurado (FOLDER_ID={bool(DRIVE_FOLDER_ID)}, CREDS={bool(CREDENTIALS_JSON)})")

    if not os.path.exists(MAPA_PATH):
        log(f"❌ Mapa não encontrado: {MAPA_PATH}")
        sys.exit(1)

    mapa = carregar_mapa()
    estado = carregar_estado()
    log(f"📊 Mapa: {mapa['total_apostilas']} | Concluídas: {len(estado.get('concluidas', []))}")

    proximas = obter_proximas_apostilas(mapa, estado, BATCH_SIZE)
    if not proximas:
        log("🎉 Todas geradas!")
        return

    log(f"📋 Próximas {len(proximas)}:")
    for apo in proximas:
        log(f"   • {apo['codigo']} — {apo['titulo']}")

    geradas = []
    falhadas = []

    for idx, meta in enumerate(proximas, 1):
        log(f"\n{'─' * 60}")
        log(f"  [{idx}/{len(proximas)}] {meta['codigo']} — {meta['titulo']}")
        log(f"  {meta['instituto']} > {meta['escola']} > {meta['curso']}")

        try:
            # 1. Gerar conteúdo
            conteudo = gerar_conteudo_gemini(meta)
            if not conteudo or len(conteudo) < 500:
                raise RuntimeError(f"Conteúdo curto: {len(conteudo) if conteudo else 0}")

            # 2. Construir docx
            os.makedirs(OUTPUT_DIR, exist_ok=True)
            fname = f"{meta['codigo']}_{slugify(meta['titulo'])}.docx"
            output_path = os.path.join(OUTPUT_DIR, fname)

            log(f"  🏗️  Construindo .docx...")
            docx_meta = {
                "titulo": meta["titulo"],
                "supratitulo": meta['instituto'],
                "subtitulo": f"{meta['escola']}  ·  Curso «{meta['curso']}»  ·  Módulo {meta['modulo']}",
                "codigo": meta["codigo"],
                "info_linha": f"Material didáctico oficial · Código {meta['codigo']} · 2026",
            }
            construir_docx(docx_meta, conteudo, output_path)
            size_kb = os.path.getsize(output_path) / 1024
            log(f"  ✅ .docx: {fname} ({size_kb:.0f} KB)")

            # 3. Upload para Google Drive
            drive_info = None
            if drive_ok:
                log(f"  ☁️  Upload para Drive...")
                try:
                    drive_info = upload_para_drive(output_path, meta)
                    if drive_info:
                        log(f"  ✅ Drive OK: {drive_info.get('webViewLink', drive_info['id'])}")
                except Exception as ue:
                    log(f"  ❌ UPLOAD FALHOU: {type(ue).__name__}: {ue}")
                    log(f"  {traceback.format_exc()}")
                    log(f"  💡 Apostila gerada localmente mas não enviada para Drive")
            else:
                log(f"  ⚠️  Upload ignorado (Drive não disponível)")

            # 4. Actualizar estado
            entry = {
                "id": meta["id"],
                "codigo": meta["codigo"],
                "titulo": meta["titulo"],
                "ficheiro": fname,
                "instituto": meta["instituto"],
                "escola": meta["escola"],
                "curso": meta["curso"],
                "modulo": meta["modulo"],
                "drive_id": drive_info['id'] if drive_info else "",
                "drive_link": drive_info.get('webViewLink', '') if drive_info else "",
            }
            geradas.append(entry)

            estado["concluidas"].append(meta["id"])
            estado["total_gerado"] = len(estado["concluidas"])
            estado.setdefault("uploads", []).append(entry)
            salvar_estado(estado)
            log(f"  💾 Estado: {estado['total_gerado']}/{mapa['total_apostilas']}")

        except Exception as e:
            log(f"  ❌ ERRO: {type(e).__name__}: {e}")
            log(f"  {traceback.format_exc()}")
            falhadas.append({"id": meta["id"], "erro": str(e)})
            estado.setdefault("falhadas", []).append(meta["id"])
            salvar_estado(estado)

        if idx < len(proximas):
            log(f"\n  ⏳ Aguardando {DELAY_ENTRE_APOSTILAS}s...")
            time.sleep(DELAY_ENTRE_APOSTILAS)

    # Resumo
    log(f"\n{'=' * 60}")
    log(f"  ✅ {len(geradas)} geradas | ❌ {len(falhadas)} falhadas")
    log(f"  Total: {len(estado.get('concluidas', []))}/{mapa['total_apostilas']}")
    uploads_ok = sum(1 for g in geradas if g.get('drive_id'))
    log(f"  ☁️  Drive: {uploads_ok}/{len(geradas)}")
    log(f"{'=' * 60}")

    if geradas:
        batch_path = os.path.join(SCRIPT_DIR, "ultimo_batch.json")
        with open(batch_path, 'w', encoding='utf-8') as f:
            json.dump(geradas, f, ensure_ascii=False, indent=2)

    if not geradas and falhadas:
        sys.exit(1)


if __name__ == "__main__":
    main()
